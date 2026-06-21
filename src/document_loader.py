from __future__ import annotations

import hashlib
from io import BytesIO
from pathlib import Path
from typing import BinaryIO

from .models import LoadedDocument


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


class DocumentLoadError(ValueError):
    """Raised when an uploaded document cannot be parsed."""


def load_uploaded_document(file_obj: BinaryIO, filename: str) -> list[LoadedDocument]:
    """Extract text from a supported uploaded document."""

    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise DocumentLoadError(f"{filename} is not supported. Use: {supported}.")

    content = file_obj.read()
    if hasattr(file_obj, "seek"):
        file_obj.seek(0)

    source_hash = hashlib.sha256(content + filename.encode("utf-8")).hexdigest()[:12]

    if suffix == ".pdf":
        return _load_pdf(content, filename, source_hash)
    if suffix == ".docx":
        return _load_docx(content, filename, source_hash)
    return _load_text(content, filename, source_hash)


def load_document_path(path: str | Path) -> list[LoadedDocument]:
    """Extract text from a document on disk while preserving source metadata."""

    document_path = Path(path)
    if not document_path.exists():
        raise DocumentLoadError(f"{document_path} does not exist.")
    if not document_path.is_file():
        raise DocumentLoadError(f"{document_path} is not a file.")

    with document_path.open("rb") as file_obj:
        return load_uploaded_document(file_obj, document_path.name)


def scan_document_folder(folder_path: str | Path, recursive: bool = True) -> list[LoadedDocument]:
    """Load every supported document from a folder."""

    folder = Path(folder_path)
    if not folder.exists():
        raise DocumentLoadError(f"{folder} does not exist.")
    if not folder.is_dir():
        raise DocumentLoadError(f"{folder} is not a folder.")

    pattern = "**/*" if recursive else "*"
    documents: list[LoadedDocument] = []
    for path in sorted(folder.glob(pattern)):
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
            documents.extend(load_document_path(path))
    return documents


def documents_to_records(documents: list[LoadedDocument]) -> list[dict[str, str | int]]:
    """Convert loaded documents into simple dictionaries for inspection or export."""

    records: list[dict[str, str | int]] = []
    for document in documents:
        record: dict[str, str | int] = {
            "text": document.text,
            "source": document.name,
        }
        record.update(document.metadata)
        records.append(record)
    return records


def _load_pdf(content: bytes, filename: str, source_hash: str) -> list[LoadedDocument]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency message for runtime.
        raise DocumentLoadError("Install pypdf to read PDF files.") from exc

    try:
        reader = PdfReader(BytesIO(content))
    except Exception as exc:  # pragma: no cover - pypdf raises multiple parser errors.
        raise DocumentLoadError(f"Could not read {filename} as a PDF.") from exc

    documents: list[LoadedDocument] = []
    for page_number, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        page_text = _clean_text(page_text)
        if not page_text:
            continue
        documents.append(
            LoadedDocument(
                source_id=f"{source_hash}-p{page_number}",
                name=filename,
                text=page_text,
                metadata={"type": "pdf", "page": page_number},
            )
        )

    if not documents:
        raise DocumentLoadError(f"No selectable text was found in {filename}.")
    return documents


def _load_docx(content: bytes, filename: str, source_hash: str) -> list[LoadedDocument]:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency message for runtime.
        raise DocumentLoadError("Install python-docx to read DOCX files.") from exc

    try:
        document = Document(BytesIO(content))
    except Exception as exc:  # pragma: no cover - python-docx parser detail.
        raise DocumentLoadError(f"Could not read {filename} as a DOCX file.") from exc

    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    text = _clean_text("\n\n".join(paragraphs))
    if not text:
        raise DocumentLoadError(f"No readable text was found in {filename}.")

    return [
        LoadedDocument(
            source_id=source_hash,
            name=filename,
            text=text,
            metadata={"type": "docx"},
        )
    ]


def _load_text(content: bytes, filename: str, source_hash: str) -> list[LoadedDocument]:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            text = content.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:  # pragma: no cover - latin-1 should always decode bytes.
        raise DocumentLoadError(f"Could not decode {filename}.")

    text = _clean_text(text)
    if not text:
        raise DocumentLoadError(f"No readable text was found in {filename}.")

    return [
        LoadedDocument(
            source_id=source_hash,
            name=filename,
            text=text,
            metadata={"type": Path(filename).suffix.lower().lstrip(".")},
        )
    ]


def _clean_text(text: str) -> str:
    lines = [line.strip() for line in text.replace("\r\n", "\n").split("\n")]
    cleaned_lines: list[str] = []
    blank_seen = False

    for line in lines:
        if not line:
            if not blank_seen:
                cleaned_lines.append("")
            blank_seen = True
            continue
        cleaned_lines.append(" ".join(line.split()))
        blank_seen = False

    return "\n".join(cleaned_lines).strip()
