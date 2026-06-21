from __future__ import annotations

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = ROOT / "data"


def main() -> None:
    DATA_DIR.mkdir(exist_ok=True)
    create_pdf(
        DATA_DIR / "business_doc.pdf",
        [
            "Business Document",
            "Northwind Analytics grew recurring revenue by 14 percent in FY2026.",
            "The company uses retrieval augmented generation to answer questions from private reports.",
            "Every answer should include citations so business users can verify the source.",
        ],
    )
    create_pdf(
        DATA_DIR / "science_paper.pdf",
        [
            "Science Paper",
            "Embeddings convert text into dense vector representations.",
            "Vector similarity search retrieves passages that are semantically close to a query.",
            "Chunk overlap preserves context when important information crosses a split boundary.",
        ],
    )
    create_docx(DATA_DIR / "factsheet.docx")
    print(f"Sample PDF and DOCX files created in {DATA_DIR}")


def create_docx(path: Path) -> None:
    document = Document()
    document.add_heading("RAG Factsheet", level=1)
    document.add_paragraph(
        "The Document Q&A Bot indexes private documents once and stores embeddings in ChromaDB."
    )
    document.add_paragraph(
        "At query time, the user question is embedded with the same model used for document chunks."
    )
    document.add_paragraph(
        "The answer prompt requires strict grounding and asks the model to cite source metadata."
    )
    document.save(path)


def create_pdf(path: Path, lines: list[str]) -> None:
    stream_lines = ["BT", "/F1 12 Tf", "72 740 Td", "16 TL"]
    for line in lines:
        stream_lines.append(f"({_escape_pdf_text(line)}) Tj")
        stream_lines.append("T*")
    stream_lines.append("ET")
    content = "\n".join(stream_lines).encode("latin-1")

    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(content)).encode("ascii") + b" >>\nstream\n" + content + b"\nendstream",
    ]

    output = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f"{index} 0 obj\n".encode("ascii"))
        output.extend(obj)
        output.extend(b"\nendobj\n")

    xref_offset = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    output.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    output.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )
    path.write_bytes(output)


def _escape_pdf_text(text: str) -> str:
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


if __name__ == "__main__":
    main()
